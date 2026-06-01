from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a Document object
doc = Document()

# Add title
title = doc.add_heading('Blog Generation Application Documentation', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add subtitle
subtitle = doc.add_paragraph('AI-Powered Content Generation with Groq LLM')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

# Add introduction
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph(
    'The Blog Generation Application is a powerful, web-based content generation tool that leverages '
    'LangChain and Groq API to automatically create high-quality blog posts. Users can specify a topic, '
    'desired word count, and target audience, and the application generates tailored blog content in seconds. '
    'The tool supports multiple writing styles optimized for different professional audiences.'
)

# Add key highlights
doc.add_heading('2. Key Highlights', 1)
doc.add_paragraph('AI-Powered Content Generation: Uses llama-3.3-70b-versatile model', style='List Bullet')
doc.add_paragraph('Multi-Audience Support: Tailored content for Researchers, Data Scientists, and Common People', style='List Bullet')
doc.add_paragraph('Customizable Length: Specify exact word count for generated blogs', style='List Bullet')
doc.add_paragraph('Web-Based Interface: No installation required, works in any browser', style='List Bullet')
doc.add_paragraph('Fast Generation: Ultra-fast inference powered by Groq', style='List Bullet')
doc.add_paragraph('Professional UI: Modern, responsive web interface with gradient design', style='List Bullet')

# Add technology stack
doc.add_heading('3. Technology Stack', 1)
doc.add_paragraph('Backend Framework: Flask - Python web server', style='List Bullet')
doc.add_paragraph('LLM Provider: Groq API with llama-3.3-70b-versatile model', style='List Bullet')
doc.add_paragraph('LangChain: PromptTemplate for structured prompt management', style='List Bullet')
doc.add_paragraph('Frontend: HTML5, CSS3, Vanilla JavaScript', style='List Bullet')
doc.add_paragraph('HTTP Client: Fetch API for asynchronous communication', style='List Bullet')
doc.add_paragraph('Environment Management: python-dotenv for API key configuration', style='List Bullet')

# Add architecture
doc.add_heading('4. System Architecture', 1)
doc.add_paragraph(
    'The application follows a client-server architecture with clear separation between '
    'frontend interface and backend processing.'
)

doc.add_heading('4.1 Backend (Flask Server)', 2)
doc.add_paragraph('HTTP Server: Runs on port 5000 for local development', style='List Bullet')
doc.add_paragraph('Route /: Serves the HTML interface', style='List Bullet')
doc.add_paragraph('Route /generate: Accepts POST requests and generates blogs', style='List Bullet')
doc.add_paragraph('LLM Integration: Manages Groq API calls for content generation', style='List Bullet')
doc.add_paragraph('Error Handling: Catches and reports errors to frontend', style='List Bullet')

doc.add_heading('4.2 Frontend (Web Interface)', 2)
doc.add_paragraph('Form Inputs: Topic, word count, and writing style selection', style='List Bullet')
doc.add_paragraph('Interactive UI: Real-time user feedback with loading indicators', style='List Bullet')
doc.add_paragraph('Async Communication: Non-blocking Fetch API requests', style='List Bullet')
doc.add_paragraph('Output Display: Scrollable blog content area', style='List Bullet')
doc.add_paragraph('Error Handling: User-friendly error messages', style='List Bullet')

# Add file structure
doc.add_heading('5. Project File Structure', 1)
doc.add_paragraph('Blog_Generation/', style='List Bullet')
doc.add_paragraph('├── app.py                    # Flask backend with embedded HTML', style='List Bullet 2')
doc.add_paragraph('└── .env                      # Environment variables (API key)', style='List Bullet 2')

# Add core components
doc.add_heading('6. Core Components', 1)

doc.add_heading('6.1 getBlogResponse Function', 2)
doc.add_paragraph('Purpose: Generate blog content using Groq LLM')
doc.add_paragraph('Parameters:', style='List Bullet')
doc.add_paragraph('input_text: Blog topic/subject matter', style='List Bullet 2')
doc.add_paragraph('no_words: Target word count for the blog', style='List Bullet 2')
doc.add_paragraph('blog_style: Target audience (Researchers, Data Scientist, Common People)', style='List Bullet 2')
doc.add_paragraph('Process:', style='List Bullet')
doc.add_paragraph('Initializes ChatGroq with llama-3.3-70b-versatile model', style='List Bullet 2')
doc.add_paragraph('Creates structured prompt template', style='List Bullet 2')
doc.add_paragraph('Invokes LLM with formatted prompt', style='List Bullet 2')
doc.add_paragraph('Returns generated blog content', style='List Bullet 2')
doc.add_paragraph('Returns: Blog content as text string', style='List Bullet')

doc.add_heading('6.2 Backend Routes', 2)

doc.add_heading('GET /', 3)
doc.add_paragraph('Purpose: Serves the main web interface')
doc.add_paragraph('Returns: HTML interface from render_template_string()')
doc.add_paragraph('URL: http://localhost:5000/')

doc.add_heading('POST /generate', 3)
doc.add_paragraph('Purpose: Processes blog generation requests')
doc.add_paragraph('Request Format:', style='List Bullet')
doc.add_paragraph('{"topic": "string", "words": "number", "style": "string"}', style='List Bullet 2')
doc.add_paragraph('Response Format:', style='List Bullet')
doc.add_paragraph('{"blog": "formatted_blog_content"}', style='List Bullet 2')
doc.add_paragraph('Validation:', style='List Bullet')
doc.add_paragraph('Ensures all fields (topic, words, style) are provided', style='List Bullet 2')
doc.add_paragraph('Returns 400 error if any field is missing', style='List Bullet 2')

doc.add_heading('6.3 Frontend Components', 2)

doc.add_heading('Form Inputs', 3)
doc.add_paragraph('Blog Topic: Text input for subject matter', style='List Bullet')
doc.add_paragraph('Number of Words: Text input for word count target', style='List Bullet')
doc.add_paragraph('Writing Style: Dropdown selector with three options', style='List Bullet')

doc.add_heading('Writing Style Options', 3)
doc.add_paragraph('Researchers: Technical, academic writing style with citations and data', style='List Bullet')
doc.add_paragraph('Data Scientist: Data-focused, analytical style with statistics', style='List Bullet')
doc.add_paragraph('Common People: Simple, accessible language for general audience', style='List Bullet')

doc.add_heading('JavaScript Functions', 3)

doc.add_heading('generateBlog(event)', 4)
doc.add_paragraph('Triggered by: Form submission or button click')
doc.add_paragraph('Actions:', style='List Bullet')
doc.add_paragraph('Prevents default form submission', style='List Bullet 2')
doc.add_paragraph('Retrieves form values (topic, words, style)', style='List Bullet 2')
doc.add_paragraph('Shows loading indicator', style='List Bullet 2')
doc.add_paragraph('Sends POST request to /generate', style='List Bullet 2')
doc.add_paragraph('Receives and displays blog content', style='List Bullet 2')
doc.add_paragraph('Handles errors with user-friendly messages', style='List Bullet 2')

# Add workflow
doc.add_heading('7. User Workflow', 1)
doc.add_paragraph('User opens application in browser (http://localhost:5000)')
doc.add_paragraph('↓')
doc.add_paragraph('User enters blog topic, word count, and selects writing style')
doc.add_paragraph('↓')
doc.add_paragraph('User clicks "Generate Blog" button')
doc.add_paragraph('↓')
doc.add_paragraph('Frontend validates inputs and sends POST request')
doc.add_paragraph('↓')
doc.add_paragraph('Backend receives request and constructs LLM prompt')
doc.add_paragraph('↓')
doc.add_paragraph('Groq API generates blog content')
doc.add_paragraph('↓')
doc.add_paragraph('Backend formats content and returns to frontend')
doc.add_paragraph('↓')
doc.add_paragraph('Frontend displays blog in scrollable container')
doc.add_paragraph('↓')
doc.add_paragraph('User can copy/save the generated content')

# Add prompt template
doc.add_heading('8. Prompt Template', 1)
doc.add_paragraph(
    'The application uses a simple yet effective prompt template that guides the LLM '
    'to generate appropriate content for the specified audience and topic.'
)
doc.add_paragraph(
    'Template: "Write a blog for {blog_style} job profile for a topic {input_text} '
    'within {no_words} words."'
)
doc.add_paragraph('This template ensures:', style='List Bullet')
doc.add_paragraph('Content is tailored for the specific audience', style='List Bullet')
doc.add_paragraph('Respects the specified word count limit', style='List Bullet')
doc.add_paragraph('Maintains professional blog writing standards', style='List Bullet')

# Add UI/UX features
doc.add_heading('9. User Interface Features', 1)

doc.add_heading('9.1 Design Elements', 2)
doc.add_paragraph('Gradient Background: Purple gradient (667eea to 764ba2) for modern look', style='List Bullet')
doc.add_paragraph('Card Layout: White container with shadow for visual hierarchy', style='List Bullet')
doc.add_paragraph('Form Design: Two-column layout for word count and style fields', style='List Bullet')
doc.add_paragraph('Color Coding: Blue focus state for inputs, purple for buttons', style='List Bullet')

doc.add_heading('9.2 Interactive Elements', 2)
doc.add_paragraph('Loading Indicator: Shows "Generating your blog..." message', style='List Bullet')
doc.add_paragraph('Button States: Disabled during processing to prevent duplicate requests', style='List Bullet')
doc.add_paragraph('Error Display: Red error messages for user awareness', style='List Bullet')
doc.add_paragraph('Output Area: Scrollable container for blog content', style='List Bullet')
doc.add_paragraph('Smooth Transitions: CSS transitions for button and input focus', style='List Bullet')

# Add configuration
doc.add_heading('10. Configuration', 1)

doc.add_heading('10.1 LLM Parameters', 2)
doc.add_paragraph('Model: llama-3.3-70b-versatile - Latest Groq model', style='List Bullet')
doc.add_paragraph('Temperature: 0.01 - Low value for consistent, factual output', style='List Bullet')
doc.add_paragraph('API Key: Retrieved from GROQ_API_KEY environment variable', style='List Bullet')

doc.add_heading('10.2 Server Configuration', 2)
doc.add_paragraph('Port: 5000 (can be changed in if __name__ == "__main__" block)', style='List Bullet')
doc.add_paragraph('Debug Mode: Enabled for development', style='List Bullet')
doc.add_paragraph('Host: localhost (127.0.0.1)', style='List Bullet')

# Add use cases
doc.add_heading('11. Real-World Use Cases', 1)
doc.add_paragraph('Content Marketing: Quickly generate blog posts for marketing campaigns', style='List Bullet')
doc.add_paragraph('Technical Documentation: Create technical blogs for developers', style='List Bullet')
doc.add_paragraph('Research Dissemination: Generate academic-style blog posts from research', style='List Bullet')
doc.add_paragraph('Product Blogs: Create blogs about new products/features', style='List Bullet')
doc.add_paragraph('SEO Content: Generate keyword-optimized blog content', style='List Bullet')
doc.add_paragraph('Educational Content: Create learning materials for students', style='List Bullet')

# Add advantages
doc.add_heading('12. Advantages', 1)
doc.add_paragraph('Fast Generation: Groq provides ultra-low latency for quick results', style='List Bullet')
doc.add_paragraph('Customizable Content: Adjustable word count and audience targeting', style='List Bullet')
doc.add_paragraph('Multiple Styles: Different writing styles for various audiences', style='List Bullet')
doc.add_paragraph('Easy to Use: Intuitive web interface requiring no technical knowledge', style='List Bullet')
doc.add_paragraph('Accessible: Works in any browser without installation', style='List Bullet')
doc.add_paragraph('Cost-Effective: Efficient token usage with Groq', style='List Bullet')

# Add limitations
doc.add_heading('13. Limitations & Considerations', 1)
doc.add_paragraph('Context Limitations: Very long prompts may hit token limits', style='List Bullet')
doc.add_paragraph('Word Count Estimation: LLM may not generate exact word count specified', style='List Bullet')
doc.add_paragraph('Content Verification: Always review generated content for accuracy', style='List Bullet')
doc.add_paragraph('API Dependency: Requires active internet and Groq API key', style='List Bullet')
doc.add_paragraph('Single User: Not configured for multi-user scenarios', style='List Bullet')

# Add improvements
doc.add_heading('14. Potential Improvements', 1)
doc.add_paragraph('Add database storage for generated blogs', style='List Bullet')
doc.add_paragraph('Implement blog export options (PDF, Markdown, HTML)', style='List Bullet')
doc.add_paragraph('Add more writing styles and tone customizations', style='List Bullet')
doc.add_paragraph('Implement user authentication and blog history', style='List Bullet')
doc.add_paragraph('Add keyword suggestion feature', style='List Bullet')
doc.add_paragraph('Implement SEO optimization suggestions', style='List Bullet')
doc.add_paragraph('Add markdown preview and editor', style='List Bullet')
doc.add_paragraph('Support for blog regeneration with different styles', style='List Bullet')

# Add setup guide
doc.add_heading('15. Setup & Deployment Guide', 1)

doc.add_heading('15.1 Prerequisites', 2)
doc.add_paragraph('Python 3.7 or higher')
doc.add_paragraph('Groq API key (from console.groq.com)')
doc.add_paragraph('pip package manager')

doc.add_heading('15.2 Installation Steps', 2)
doc.add_paragraph('1. Create project directory:', style='List Number')
doc.add_paragraph('mkdir Blog_Generation && cd Blog_Generation', style='List Number 2')
doc.add_paragraph('2. Create and activate virtual environment:', style='List Number')
doc.add_paragraph('python -m venv venv', style='List Number 2')
doc.add_paragraph('venv\\Scripts\\activate  (Windows)', style='List Number 2')
doc.add_paragraph('3. Install dependencies:', style='List Number')
doc.add_paragraph('pip install flask langchain-groq langchain-core python-dotenv', style='List Number 2')
doc.add_paragraph('4. Create .env file with API key:', style='List Number')
doc.add_paragraph('echo GROQ_API_KEY=your_key_here > .env', style='List Number 2')
doc.add_paragraph('5. Run application:', style='List Number')
doc.add_paragraph('python app.py', style='List Number 2')

doc.add_heading('15.3 Running Locally', 2)
doc.add_paragraph('Start the server: python app.py')
doc.add_paragraph('Open browser: http://localhost:5000')
doc.add_paragraph('Begin generating blogs!')

# Add example usage
doc.add_heading('16. Example Usage Scenarios', 1)

doc.add_heading('Example 1: Technical Blog for Data Scientists', 2)
doc.add_paragraph('Topic: "Machine Learning in Healthcare"')
doc.add_paragraph('Words: 800')
doc.add_paragraph('Style: Data Scientist')
doc.add_paragraph('Result: Technical blog with statistics, algorithms, and use cases')

doc.add_heading('Example 2: Academic Blog for Researchers', 2)
doc.add_paragraph('Topic: "Quantum Computing Breakthroughs"')
doc.add_paragraph('Words: 1200')
doc.add_paragraph('Style: Researchers')
doc.add_paragraph('Result: Research-focused blog with citations and methodologies')

doc.add_heading('Example 3: General Audience Blog', 2)
doc.add_paragraph('Topic: "How Artificial Intelligence Affects Daily Life"')
doc.add_paragraph('Words: 500')
doc.add_paragraph('Style: Common People')
doc.add_paragraph('Result: Accessible, easy-to-understand blog for general readers')

# Add security notes
doc.add_heading('17. Security Best Practices', 1)
doc.add_paragraph('Never commit .env file to version control', style='List Bullet')
doc.add_paragraph('Keep GROQ_API_KEY confidential and secure', style='List Bullet')
doc.add_paragraph('Validate all user inputs on frontend and backend', style='List Bullet')
doc.add_paragraph('For production: Implement rate limiting', style='List Bullet')
doc.add_paragraph('For production: Add HTTPS/SSL certificate', style='List Bullet')
doc.add_paragraph('Monitor API usage to prevent unexpected charges', style='List Bullet')

# Add troubleshooting
doc.add_heading('18. Troubleshooting', 1)

doc.add_heading('Issue: "GROQ_API_KEY not found"', 2)
doc.add_paragraph('Solution: Ensure .env file exists in project root with correct key')

doc.add_heading('Issue: "Connection refused on port 5000"', 2)
doc.add_paragraph('Solution: Check if port 5000 is already in use, or change port in app.py')

doc.add_heading('Issue: "Blog generation taking too long"', 2)
doc.add_paragraph('Solution: Check internet connection and Groq API status')

doc.add_heading('Issue: "Word count not matching"', 2)
doc.add_paragraph('Solution: LLM may approximate - provide specific word targets as guidelines')

# Add conclusion
doc.add_heading('19. Conclusion', 1)
doc.add_paragraph(
    'The Blog Generation Application demonstrates the power of combining LangChain\'s prompt '
    'management with Groq\'s high-performance LLM inference. By providing an intuitive web interface '
    'with support for multiple audience types, this tool enables rapid content creation for various use cases. '
    'Whether for marketing, technical documentation, or educational content, this application provides a '
    'solid foundation for AI-powered blog generation. The architecture is simple yet scalable, making it '
    'easy to extend with additional features and customizations.'
)

# Save the document
output_path = r'c:\Users\KavyaS\OneDrive - Preferred Square Analytics Pvt Ltd\Documents\Learn\Projects\prompt + API\Blog_Generation\Blog_Generation_Documentation.docx'
doc.save(output_path)
print(f'Blog Generation documentation created successfully at: {output_path}')
