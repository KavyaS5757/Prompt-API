from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a Document object
doc = Document()

# Add title
title = doc.add_heading('RAG Project Documentation', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Retrieval-Augmented Generation with PDF Processing')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

# Add introduction
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph(
    'This RAG (Retrieval-Augmented Generation) project is a sophisticated document processing and '
    'question-answering system that combines PDF text extraction, semantic search, and advanced '
    'language models. Users can upload PDF documents and ask questions about them, receiving '
    'accurate, context-aware answers powered by the Groq API and Chroma vector database.'
)

# Add what is RAG
doc.add_heading('2. What is RAG?', 1)
doc.add_paragraph(
    'RAG (Retrieval-Augmented Generation) is an AI technique that combines two key steps:'
)
doc.add_paragraph(
    'Retrieval: Finding relevant documents or text chunks based on user queries using semantic search',
    style='List Number'
)
doc.add_paragraph(
    'Generation: Using an LLM to generate accurate answers based on the retrieved context',
    style='List Number'
)
doc.add_paragraph(
    'This approach ensures answers are grounded in actual document content rather than relying '
    'solely on the LLM\'s training data, improving accuracy and reducing hallucinations.'
)

# Add architecture
doc.add_heading('3. System Architecture', 1)
doc.add_paragraph(
    'The project consists of three main Python modules working together:'
)

doc.add_heading('3.1 ingest.py - Document Ingestion Pipeline', 2)
doc.add_paragraph(
    'This module handles PDF processing and storage in the vector database:'
)
doc.add_paragraph('PDF Reading: Uses PyMuPDF (fitz) to extract text from PDF files', style='List Bullet')
doc.add_paragraph('Text Chunking: Splits documents into 500-word chunks with 50-word overlap', style='List Bullet')
doc.add_paragraph('Embedding: Converts text chunks to vector embeddings using SentenceTransformer', style='List Bullet')
doc.add_paragraph('Storage: Stores embeddings and chunks in Chroma persistent database', style='List Bullet')
doc.add_paragraph('Cleanup: Removes old versions of documents to avoid duplicates', style='List Bullet')

doc.add_heading('Key Functions:', 3)
doc.add_paragraph('load_pdf(filepath): Extracts all text from a PDF document', style='List Bullet')
doc.add_paragraph('chunk_text(text, chunk_size, overlap): Splits text into overlapping chunks', style='List Bullet')
doc.add_paragraph('ingest_pdf(filepath): Main pipeline orchestrating the entire ingestion process', style='List Bullet')

doc.add_heading('3.2 rag.py - Retrieval and Answer Generation', 2)
doc.add_paragraph(
    'This module implements the core RAG functionality:'
)
doc.add_paragraph('Vector Search: Embeds user questions and finds relevant document chunks', style='List Bullet')
doc.add_paragraph('Context Building: Assembles retrieved chunks into a coherent context', style='List Bullet')
doc.add_paragraph('LLM Processing: Sends context and question to Groq for answer generation', style='List Bullet')
doc.add_paragraph('Response Formatting: Returns both the answer and source chunks', style='List Bullet')

doc.add_heading('Key Functions:', 3)
doc.add_paragraph('get_collections(): Lists all available documents in the database', style='List Bullet')
doc.add_paragraph('ask(question, collection_name, n_results): Queries a document and retrieves answers', style='List Bullet')

doc.add_heading('3.3 app.py - Web Interface', 2)
doc.add_paragraph(
    'This Flask application provides a user-friendly web interface:'
)
doc.add_paragraph('PDF Upload: Allows users to upload documents via file picker', style='List Bullet')
doc.add_paragraph('Document Selection: Dropdown to choose which document to query', style='List Bullet')
doc.add_paragraph('Chat Interface: Interactive Q&A interface with message history', style='List Bullet')
doc.add_paragraph('Source Tracking: Shows which document chunks support each answer', style='List Bullet')

# Add technology stack
doc.add_heading('4. Technology Stack', 1)
doc.add_paragraph('Framework: Flask - Lightweight Python web framework', style='List Bullet')
doc.add_paragraph('PDF Processing: PyMuPDF (fitz) - Fast PDF text extraction', style='List Bullet')
doc.add_paragraph('Vector Database: Chroma - Open-source embeddings database', style='List Bullet')
doc.add_paragraph('Embeddings: SentenceTransformer (all-MiniLM-L6-v2) - Fast, free embeddings', style='List Bullet')
doc.add_paragraph('LLM: Groq API with llama-3.3-70b-versatile - Ultra-fast inference', style='List Bullet')
doc.add_paragraph('Frontend: HTML/CSS/JavaScript - Interactive web UI', style='List Bullet')
doc.add_paragraph('Environment: python-dotenv - Secure API key management', style='List Bullet')

# Add workflow
doc.add_heading('5. Workflow Process', 1)

doc.add_heading('5.1 Document Upload & Ingestion', 2)
doc.add_paragraph('User uploads a PDF file through the web interface')
doc.add_paragraph('↓')
doc.add_paragraph('Flask receives the file and saves it to uploads/ folder')
doc.add_paragraph('↓')
doc.add_paragraph('ingest.py extracts text using PyMuPDF')
doc.add_paragraph('↓')
doc.add_paragraph('Text is split into overlapping 500-word chunks')
doc.add_paragraph('↓')
doc.add_paragraph('SentenceTransformer converts chunks to vector embeddings')
doc.add_paragraph('↓')
doc.add_paragraph('Chroma stores embeddings with document references')
doc.add_paragraph('↓')
doc.add_paragraph('Collection appears in document selector')

doc.add_heading('5.2 Question Answering', 2)
doc.add_paragraph('User types a question and selects a document')
doc.add_paragraph('↓')
doc.add_paragraph('Question is converted to embedding by SentenceTransformer')
doc.add_paragraph('↓')
doc.add_paragraph('Chroma performs semantic similarity search to find top-4 relevant chunks')
doc.add_paragraph('↓')
doc.add_paragraph('Retrieved chunks are combined into context string')
doc.add_paragraph('↓')
doc.add_paragraph('Context and question sent to Groq LLM')
doc.add_paragraph('↓')
doc.add_paragraph('Groq generates answer based on document context')
doc.add_paragraph('↓')
doc.add_paragraph('Answer displayed with source chunk references')

# Add key features
doc.add_heading('6. Key Features', 1)
doc.add_paragraph('Multiple Document Support: Manage and query multiple PDFs simultaneously', style='List Bullet')
doc.add_paragraph('Semantic Search: Finds relevant chunks by meaning, not just keywords', style='List Bullet')
doc.add_paragraph('Source Attribution: Shows which parts of the document support each answer', style='List Bullet')
doc.add_paragraph('Fast Inference: Groq provides ultra-fast LLM responses', style='List Bullet')
doc.add_paragraph('Overlap Chunking: Preserves context between chunks for better understanding', style='List Bullet')
doc.add_paragraph('Factual Accuracy: Low temperature setting ensures consistent, grounded responses', style='List Bullet')
doc.add_paragraph('Clean UI: Intuitive web interface with real-time chat experience', style='List Bullet')

# Add configuration parameters
doc.add_heading('7. Configuration Parameters', 1)

doc.add_heading('7.1 Chunking Configuration', 2)
doc.add_paragraph('chunk_size=500: Number of words per chunk', style='List Bullet')
doc.add_paragraph('overlap=50: Words overlapped between adjacent chunks for context continuity', style='List Bullet')
doc.add_paragraph(
    'Note: Adjust these based on document type - smaller chunks for technical docs, '
    'larger for narrative text',
    style='List Bullet'
)

doc.add_heading('7.2 Search Configuration', 2)
doc.add_paragraph('n_results=4: Number of similar chunks retrieved per query', style='List Bullet')
doc.add_paragraph('temperature=0.2: Low value ensures factual, consistent answers', style='List Bullet')
doc.add_paragraph('Embeddings Model: all-MiniLM-L6-v2 - Fast semantic embeddings', style='List Bullet')

# Add use cases
doc.add_heading('8. Real-World Use Cases', 1)
doc.add_paragraph('Legal Document Analysis: Review contracts and extract specific clauses', style='List Bullet')
doc.add_paragraph('Research Paper Search: Find specific findings across multiple papers', style='List Bullet')
doc.add_paragraph('Customer Support: Answer questions based on product documentation', style='List Bullet')
doc.add_paragraph('Knowledge Base Search: Query company policies and procedures', style='List Bullet')
doc.add_paragraph('Academic Assistance: Understand complex textbooks and research materials', style='List Bullet')
doc.add_paragraph('Financial Analysis: Extract insights from financial reports', style='List Bullet')

# Add benefits
doc.add_heading('9. Advantages Over Traditional Approaches', 1)
doc.add_paragraph('Higher Accuracy: Answers grounded in actual documents', style='List Bullet')
doc.add_paragraph('Reduced Hallucinations: LLM can\'t invent information not in documents', style='List Bullet')
doc.add_paragraph('Source Attribution: Users can verify answers against source material', style='List Bullet')
doc.add_paragraph('Scalability: Can handle thousands of documents with persistent storage', style='List Bullet')
doc.add_paragraph('Cost-Effective: Reusable embeddings avoid re-processing documents', style='List Bullet')
doc.add_paragraph('Fast Performance: Groq\'s optimization enables real-time responses', style='List Bullet')

# Add limitations and considerations
doc.add_heading('10. Limitations & Considerations', 1)
doc.add_paragraph('PDF Quality: Scanned PDFs without OCR won\'t extract text properly', style='List Bullet')
doc.add_paragraph('Context Size: Limited to top-4 chunks; may miss relevant information in long documents', style='List Bullet')
doc.add_paragraph('Embedding Limitations: SentenceTransformer optimized for semantic meaning, not technical terms', style='List Bullet')
doc.add_paragraph('Language: Works best with English documents', style='List Bullet')
doc.add_paragraph('API Dependencies: Requires active Groq API key and internet connection', style='List Bullet')

# Add future improvements
doc.add_heading('11. Potential Improvements', 1)
doc.add_paragraph('Implement hierarchical chunking for better context preservation', style='List Bullet')
doc.add_paragraph('Add OCR support for scanned PDFs', style='List Bullet')
doc.add_paragraph('Support multiple document formats (DOCX, TXT, markdown)', style='List Bullet')
doc.add_paragraph('Implement question rephrasing for better chunk retrieval', style='List Bullet')
doc.add_paragraph('Add document summarization features', style='List Bullet')
doc.add_paragraph('Implement conversation history for follow-up questions', style='List Bullet')
doc.add_paragraph('Add user authentication and document privacy controls', style='List Bullet')

# Add example scenarios
doc.add_heading('12. Example Usage Scenarios', 1)

doc.add_heading('Scenario 1: Legal Contract Analysis', 2)
doc.add_paragraph('User: "What are the termination conditions in this contract?"')
doc.add_paragraph('System: Searches for relevant clauses and extracts termination terms')
doc.add_paragraph('Result: Returns specific contract sections with direct quotes')

doc.add_heading('Scenario 2: Research Paper Query', 2)
doc.add_paragraph('User: "What methodology was used in this study?"')
doc.add_paragraph('System: Retrieves methodology section chunks and synthesizes answer')
doc.add_paragraph('Result: Provides clear methodology explanation with source citations')

doc.add_heading('Scenario 3: Policy Question', 2)
doc.add_paragraph('User: "What is the vacation policy for employees?"')
doc.add_paragraph('System: Searches employee handbook for vacation-related sections')
doc.add_paragraph('Result: Returns vacation policy details with relevant policy document chunks')

# Add environment setup
doc.add_heading('13. Environment Setup', 1)
doc.add_paragraph('Required environment variable:')
doc.add_paragraph('GROQ_API_KEY: Your Groq API key from console.groq.com', style='List Bullet')
doc.add_paragraph('Store in .env file in the project root')

# Add conclusion
doc.add_heading('14. Conclusion', 1)
doc.add_paragraph(
    'The RAG project demonstrates a production-ready document question-answering system that '
    'combines modern NLP techniques with fast inference. By integrating semantic search with '
    'large language models, it provides accurate, attributed answers to questions about PDF documents. '
    'This solution can be extended for enterprise document management, customer support automation, '
    'and knowledge base applications.'
)

# Save the document
output_path = r'c:\Users\KavyaS\OneDrive - Preferred Square Analytics Pvt Ltd\Documents\Learn\Projects\prompt + API\rag_project\RAG_Project_Documentation.docx'
doc.save(output_path)
print(f'RAG Project documentation created successfully at: {output_path}')
