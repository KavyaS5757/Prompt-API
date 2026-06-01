from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a Document object
doc = Document()

# Add title
title = doc.add_heading('Langchain with Groq API - Use Case Documentation', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add introduction
doc.add_heading('1. Overview', 1)
doc.add_paragraph(
    'This Flask application demonstrates a sophisticated multi-step LLM chain that combines '
    'the power of LangChain and Groq API to create an intelligent information retrieval system. '
    'The application processes user queries about celebrities and returns comprehensive information '
    'in a structured format.'
)

# Add main usecase
doc.add_heading('2. Primary Use Case', 1)
doc.add_paragraph(
    'Celebrity Information Research & Timeline Analysis'
)
doc.add_paragraph(
    'The application allows users to search for a celebrity\'s name and automatically retrieves:'
)
doc.add_paragraph('Detailed biographical information about the celebrity', style='List Bullet')
doc.add_paragraph('Their date of birth', style='List Bullet')
doc.add_paragraph('Major world events that occurred around their birth year', style='List Bullet')

# Add how it works
doc.add_heading('3. How It Works', 1)
doc.add_paragraph('The application uses a 3-step LLM chain:')

doc.add_heading('Step 1: Celebrity Information Retrieval', 2)
doc.add_paragraph(
    'Template: "Tell me about celebrity {name}"'
)
doc.add_paragraph(
    'The first prompt uses the user\'s input to fetch comprehensive information about the celebrity '
    'using the Groq LLM.'
)

doc.add_heading('Step 2: Birth Date Extraction', 2)
doc.add_paragraph(
    'Template: "when was {person} born"'
)
doc.add_paragraph(
    'The second prompt takes the celebrity information from Step 1 and extracts their birth date. '
    'This demonstrates prompt chaining where output from one step feeds into the next.'
)

doc.add_heading('Step 3: Historical Events Analysis', 2)
doc.add_paragraph(
    'Template: "Mention 5 major events happened around {dob} in the world"'
)
doc.add_paragraph(
    'The final prompt uses the birth date from Step 2 to generate a list of 5 major world events '
    'that occurred around that time period, providing historical context.'
)

# Add technical details
doc.add_heading('4. Technical Stack', 1)
doc.add_paragraph('Framework: Flask - Lightweight web framework for Python', style='List Bullet')
doc.add_paragraph('LLM Provider: Groq API with llama-3.3-70b-versatile model', style='List Bullet')
doc.add_paragraph('LangChain: PromptTemplate for structured prompt management', style='List Bullet')
doc.add_paragraph('Frontend: HTML template for user interface', style='List Bullet')

# Add key features
doc.add_heading('5. Key Features', 1)
doc.add_paragraph('Multi-Step Prompt Chaining: Demonstrates how to chain multiple LLM calls', style='List Bullet')
doc.add_paragraph('Environment Variable Management: Uses .env file for secure API key storage', style='List Bullet')
doc.add_paragraph('Error Handling: Catches and displays errors gracefully', style='List Bullet')
doc.add_paragraph('Template-Based Prompting: Uses LangChain PromptTemplate for maintainability', style='List Bullet')
doc.add_paragraph('Web Interface: User-friendly Flask web application', style='List Bullet')

# Add real-world applications
doc.add_heading('6. Real-World Applications', 1)
doc.add_paragraph('Educational Platforms: History lessons with biographical and contextual information', style='List Bullet')
doc.add_paragraph('Research Tools: Automated biography and timeline generation', style='List Bullet')
doc.add_paragraph('Content Creation: Blog post or article generation about notable figures', style='List Bullet')
doc.add_paragraph('Knowledge Base: Building structured information from unstructured queries', style='List Bullet')
doc.add_paragraph('Chatbots: Advanced Q&A systems with multi-step reasoning', style='List Bullet')

# Add benefits
doc.add_heading('7. Benefits of This Approach', 1)
doc.add_paragraph('High Performance: Groq provides ultra-fast LLM inference', style='List Bullet')
doc.add_paragraph('Structured Output: PromptTemplate ensures consistent prompt formatting', style='List Bullet')
doc.add_paragraph('Scalability: Easy to extend with more chains and prompts', style='List Bullet')
doc.add_paragraph('Maintainability: Clear separation of concerns with Flask routes and LangChain templates', style='List Bullet')
doc.add_paragraph('Cost-Effective: Efficient processing with Groq\'s optimized infrastructure', style='List Bullet')

# Add example workflow
doc.add_heading('8. Example Workflow', 1)
doc.add_paragraph('User Input: "Albert Einstein"')
doc.add_paragraph('↓')
doc.add_paragraph('Step 1 Output: Detailed biography of Albert Einstein')
doc.add_paragraph('↓')
doc.add_paragraph('Step 2 Output: Born on March 14, 1879')
doc.add_paragraph('↓')
doc.add_paragraph('Step 3 Output: 5 major world events from 1879 (Industrial Revolution peak, etc.)')
doc.add_paragraph('↓')
doc.add_paragraph('Final Result: Comprehensive profile displayed to user')

# Add conclusion
doc.add_heading('9. Conclusion', 1)
doc.add_paragraph(
    'This application showcases the power of combining LangChain\'s prompt management capabilities '
    'with Groq\'s fast LLM inference to create sophisticated, multi-step AI workflows. It demonstrates '
    'practical prompt engineering, information retrieval, and web application development, making it '
    'an excellent example for learning advanced LLM integration techniques.'
)

# Save the document
output_path = r'c:\Users\KavyaS\OneDrive - Preferred Square Analytics Pvt Ltd\Documents\Learn\Projects\prompt + API\Langchain\Usecase_Documentation.docx'
doc.save(output_path)
print(f'Document created successfully at: {output_path}')
